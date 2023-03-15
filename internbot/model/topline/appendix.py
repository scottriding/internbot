from openpyxl import load_workbook, Workbook
from openpyxl.styles.borders import Border, Side
from openpyxl.styles import PatternFill, Font, Alignment
from openpyxl.drawing.image import Image
import os
import csv
from collections import OrderedDict
import docx

class Appendix(object):

	def check_input(self, csv_path):
		self.check_input_names(open(csv_path, encoding="utf8"))

	def check_input_names(self, utf8_data, **kwargs):
		csv_reader = csv.DictReader(utf8_data, **kwargs)
		for row in csv_reader:
			if row.get("variable") is None:
				raise ValueError(f'Missing column: variable')
			if row.get("prompt") is None:
				raise ValueError(f'Missing column: prompt')
			if row.get("label") is None:
				raise ValueError(f'Missing column: label')
			break

	def build_appendix_report(self, question_blocks, path_to_output, template_path):
		builder = Document(template_path)

		builder.write_appendix(question_blocks)
		builder.save(path_to_output)
		self.__questions = OrderedDict() ## empty out questions

class Document(object):

	def __init__ (self, path_to_template):
		self.__doc = docx.Document(path_to_template)

	def write_appendix(self, question_blocks):
		first_question = True
		for question in question_blocks.questions:
			if first_question is False:
				self.__doc.add_page_break()

			paragraph = self.__doc.add_paragraph()
			self.write_question(question, paragraph)
			self.__doc.add_paragraph()
			first_question = False


	def write_question(self, question, paragraph):
		paragraph.add_run(question.name + ".")
		paragraph_format = paragraph.paragraph_format
		paragraph_format.keep_together = True
		paragraph_format.left_indent = docx.shared.Inches(1)
		prompt_to_add = "\t%s (n = %s)\n" % (question.prompt, len(question.responses))
		paragraph.add_run(prompt_to_add)
		paragraph_format.first_line_indent = docx.shared.Inches(-1)

		self.write_responses(question.responses, paragraph)


	def write_responses(self, responses, paragraph):
		responses.appendix_sort()
		table = self.__doc.add_table(rows = 0, cols = 5)
		try:
			table.style = 'Appendix' # Custom format in template
		except KeyError:
			pass
		for response in responses:
			response_cells = table.add_row().cells
			response_cells[0].merge(response_cells[4])
			response_cells[0].text = response.label

	def save(self, path_to_output):
		self.__doc.save(path_to_output)

