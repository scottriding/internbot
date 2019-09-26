from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

class CSVToplineReport(object):

    def __init__(self, questions, path_to_template, years):
        self.doc = Document(path_to_template)
        self.line_break = self.doc.styles['LineBreak']
        self.questions = questions
        self.years = years
        names = questions.list_names()
        for name in names:
            self.write_question(name)
        print("Finished!")
        
    def write_question(self, name):
        question = self.questions.get(name)
        paragraph = self.doc.add_paragraph() # each question starts a new paragraph
        self.write_name(name, paragraph)
        self.write_prompt(question.prompt, paragraph)
        self.write_n(question.n, paragraph)
        if len(question.responses) > 0:
            self.write_responses(question.responses, question.stat)
        new = self.doc.add_paragraph("") # space between questions
        new.style = self.line_break
        self.doc.add_paragraph("") # space between questions
        
    def write_name(self, name, paragraph):
        paragraph.add_run(name + ".")
    
    def write_prompt(self, prompt, paragraph):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.keep_together = True           # question prompt will all be fit in one page
        paragraph_format.left_indent = Inches(1)        
        paragraph.add_run("\t" + prompt)                
        paragraph_format.first_line_indent = Inches(-1) # hanging indent if necessary
        
    def write_n(self, n, paragraph):
        paragraph.add_run(" (n = " + str(n) + ")")
    
    def write_responses(self, responses, stat):
        if len(self.years) > 0:
            self.write_trended_responses(responses, stat)
        else:
            table = self.doc.add_table(rows = 1, cols = 5)
            first_row = True
            for response in responses:
                response_cells = table.add_row().cells
                response_cells[1].merge(response_cells[2])
                response_cells[1].text = response.name
                for year, response in response.frequencies.items():
                    if stat == 'percent':
                        response_cells[3].text = self.freqs_percent(response, first_row)
                    else:
                        response_cells[3].text = str(response)
                if response_cells[3].text != "*":
                    first_row = False

    def write_trended_responses(self, responses, stat):
        headers =  self.max_years(responses)
        table = self.doc.add_table(rows=1, cols=len(headers)+4)
        titles_row = table.add_row().cells  
        titles_row[1].merge(titles_row[2])
        headers_index = 0
        while headers_index < len(headers):
            header_text = "Total %s" % headers[headers_index]
            titles_row[headers_index+4].text = header_text
            headers_index += 1
        first_row = True
        for response in responses:
            response_cells = table.add_row().cells
            response_cells[1].merge(response_cells[3])
            response_cells[1].text = response.name
            freq_col = 4
            for header in headers:
                if response.frequencies.get(header) is not None:
                    freq = response.frequencies.get(header)
                    if stat == 'percent':
                        text = self.freqs_percent(freq, first_row)
                    else:
                        text = str(freq)
                    response_cells[freq_col].text = text
                first_row = False
                freq_col += 1

    def max_years(self, responses):
        years_used = []
        for response in responses:
            for year in self.years:
                if response.frequencies.get(year) is not None:
                    if year not in years_used:
                        years_used.append(year)
        return years_used

    def save(self, path_to_output):
        self.doc.save(path_to_output)
        
    def freqs_percent(self, freq, is_first=False):
        result = 0
        if float(freq) >= 1.0:
            result = int(freq) * 100
        else:
            percent = float(freq)
            percent = percent * 100
            if percent > 0 and percent < 1:
                result = "<1"
            elif percent == 0:
                result = "*"
            else:
                result = int(round(percent))
        if is_first and result != "*":
            result = str(result) + "%"
        return str(result)
