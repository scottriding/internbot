from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class QSFToplineReport(object):

    def __init__ (self, questions, path_to_template, years):
        self.doc = Document(path_to_template)
        self.line_break = self.doc.styles['LineBreak']
        self.questions = questions
        self.years = years

    def save(self, path_to_output):
        self.write_questions()
        self.save_file(path_to_output)
        print("Finished!")

    def write_questions(self):
        for question in self.questions:
            to_print = "Writing question: %s" % question.name
            print(to_print)

            if question.parent == 'CompositeQuestion':
                self.write_composite_question(question)
            elif question.type == 'TE':
                self.write_open_ended(question)
            else:
                self.write_question(question)

    def save_file(self, path_to_output):
        self.doc.save(path_to_output)

    def write_question(self, question):
        #paragraph = self.doc.add_paragraph() # each question starts a new paragraph
        #self.write_display_logic(question, paragraph)
        paragraph = self.doc.add_paragraph()
        self.write_name(question.name, paragraph)
        self.write_prompt(question.prompt, paragraph)
        self.write_n(question.n, paragraph)
        if question.type == 'RO':
            self.write_rank(question.responses)
        else:
            self.write_responses(question.responses, question.stat)
        new = self.doc.add_paragraph("") # space between questions
        new.style = self.line_break
        self.doc.add_paragraph("") # space between questions

    def write_composite_question(self, question):
        #paragraph = self.doc.add_paragraph() # each question starts a new paragraph
        #self.write_display_logic(question, paragraph)
        paragraph = self.doc.add_paragraph()
        self.write_name(question.name, paragraph)
        self.write_prompt(question.prompt, paragraph)
        if question.type == 'CompositeMatrix':
            self.write_matrix(question.questions)
        elif question.type == 'CompositeConstantSum':
            self.write_allocate(question.questions)
        else:
            self.write_binary(question.questions)
        new = self.doc.add_paragraph("") # space between questions
        new.style = self.line_break
        self.doc.add_paragraph("")

    def write_open_ended(self, question):
        paragraph = self.doc.add_paragraph()  # each question starts a new paragraph
        self.write_name(question.name, paragraph)
        self.write_prompt(question.prompt, paragraph)
        self.write_n(question.n, paragraph)
        paragraph.add_run(' (OPEN-ENDED RESPONSES VERBATIM IN APPENDIX)')
        new = self.doc.add_paragraph("") # space between questions
        new.style = self.line_break
        self.doc.add_paragraph("")

    def write_name(self, name, paragraph):
        paragraph.add_run(name + ".")

    def write_prompt(self, prompt, paragraph):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.keep_together = True # question prompt will all be fit in one page
        paragraph_format.left_indent = Inches(1)
        paragraph.add_run("\t" + prompt)
        paragraph_format.first_line_indent = Inches(-1) # hanging indent if necessary

    def write_n(self, n, paragraph):
        if n != 0:
            paragraph.add_run(" (n = " + str(n) + ")")

    def write_display_logic(self, question, paragraph):
        # This is broken still needs to be implemented
        if question.parent != 'CompositeQuestion' and question.type != 'CompositeMultipleSelect' and question.display_logic != "" and question.display_logic is not None:
            paragraph.add_run(str(question.display_logic))
        elif question.parent == 'CompositeQuestion':
            if question.type != 'CompositeMultipleSelect':
                for sub in question.questions:
                    if sub.display_logic != "":
                        question.display_logic = sub.display_logic
                        break
                paragraph.add_run(str(question.display_logic))


    def write_responses(self, responses, stat):
        if len(self.years) > 0:
            self.write_trended_responses(responses, stat)
        else:
            table = self.doc.add_table(rows=1, cols=5)
            first_row = True
            for response in responses:
                response_cells = table.add_row().cells
                response_cells[1].merge(response_cells[2])
                response_cells[1].text = response.response
                if not response.frequencies:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="FFF206"/>'.format(nsdecls('w')))
                    response_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
                for year, response in response.frequencies.items():
                    if stat == 'percent':
                        freq = self.freqs_percent(response, first_row)
                    else:
                        freq = str(response)
                    response_cells[3].text = freq
                first_row = False

    def write_trended_responses(self, responses, stat):
        headers = self.max_years(responses)
        table = self.doc.add_table(rows=1, cols=len(headers) + 4)
        titles_row = table.add_row().cells
        titles_row[1].merge(titles_row[2])
        headers_index = 0
        while headers_index < len(headers):
            header_text = "Total %s" % headers[headers_index]
            titles_row[headers_index + 4].text = header_text
            headers_index += 1
        first_row = True
        for response in responses:
            response_cells = table.add_row().cells
            response_cells[1].merge(response_cells[3])
            response_cells[1].text = response.response
            freq_col = 4
            if not response.frequencies:
                shading_elm = parse_xml(r'<w:shd {} w:fill="FFF206"/>'.format(nsdecls('w')))
                response_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
            for header in headers:
                if response.frequencies.get(header) is not None:
                    freq = response.frequencies.get(header)
                    if stat == 'percent':
                        text = self.freqs_percent(freq, first_row)
                    else:
                        text = str(freq)
                    response_cells[freq_col].text = text
                first_row = False
                freq_col += 1

    def write_rank(self, responses):
        if len(self.years) > 0:
            self.write_trended_rank(responses)
        else:
            table = self.doc.add_table(rows=1, cols=5)
            first_row = True
            for response in responses:
                response_cells = table.add_row().cells
                response_cells[1].merge(response_cells[2])
                response_cells[1].text = response.response
                if not response.frequencies:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="FFF206"/>'.format(nsdecls('w')))
                    response_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
                for year, average in response.frequencies.items():
                    response_cells[3].text = self.avg_float(average, first_row)
                first_row = False

    def write_trended_rank(self, responses):
        headers = self.max_years(responses)
        table = self.doc.add_table(rows=1, cols=len(headers) + 4)
        titles_row = table.add_row().cells
        titles_row[1].merge(titles_row[2])
        headers_index = 0
        while headers_index < len(headers):
            header_text = "Total %s" % headers[headers_index]
            titles_row[headers_index + 4].text = header_text
            headers_index += 1
        first_row = True
        for response in responses:
            response_cells = table.add_row().cells
            response_cells[1].merge(response_cells[3])
            response_cells[1].text = response.response
            freq_col = 4
            if not response.frequencies:
                shading_elm = parse_xml(r'<w:shd {} w:fill="FFF206"/>'.format(nsdecls('w')))
                response_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
            for header in headers:
                if response.frequencies.get(header) is not None:
                    avg = response.frequencies.get(header)
                    text = self.avg_float(avg, first_row)
                    response_cells[freq_col].text = text
                first_row = False
                freq_col += 1

    def write_binary(self, sub_questions):
        if len(self.years) > 0:
            self.write_trended_binary(sub_questions)
        else:
            table = self.doc.add_table(rows=1, cols=5)
            first_row = True
            for sub_question in sub_questions:
                cells = table.add_row().cells
                cells[1].merge(cells[2])
                cells[1].text = "%s (n=%s)" % (sub_question.prompt, sub_question.n)
                response = next((response for response in sub_question.responses if response.code == '1'), None)
                if not response.frequencies:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="FFF206"/>'.format(nsdecls('w')))
                    cells[1]._tc.get_or_add_tcPr().append(shading_elm)
                for year, frequency in response.frequencies.items():
                    cells[3].text = self.freqs_percent(frequency, first_row)
                first_row = False

    def write_trended_binary(self, sub_questions):
        headers = self.max_years_subquestions(sub_questions)
        table = self.doc.add_table(rows=1, cols=len(headers)+4)
        titles_row = table.add_row().cells
        titles_row[1].merge(titles_row[2])
        headers_index = 0
        while headers_index < len(headers):
            header_text = "Total %s" % headers[headers_index]
            titles_row[headers_index+4].text = header_text
            headers_index += 1
        first_row = True
        for sub_question in sub_questions:
            response = next((response for response in sub_question.responses if response.code == '1'), None)
            region_cells = table.add_row().cells
            region_cells[1].merge(region_cells[3])
            region_cells[1].text = "%s (n=%s)" % (sub_question.prompt, sub_question.n)
            freq_col = 4
            if not response.frequencies:
                shading_elm = parse_xml(r'<w:shd {} w:fill="FFF206"/>'.format(nsdecls('w')))
                region_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
            for header in headers:
                if response.frequencies.get(header) is not None:
                    freq = response.frequencies.get(header)
                    text = self.freqs_percent(freq, first_row)
                    region_cells[freq_col].text = text
                else:
                    if first_row is True:
                        region_cells[freq_col].text = "--%"
                    else:
                        region_cells[freq_col].text = "--"
                first_row = False
                freq_col += 1

    def write_allocate(self, sub_questions):
        if len(self.years) > 0:
            self.write_trended_allocate(sub_questions)
        else:
            table = self.doc.add_table(rows=1, cols=5)
            first_row = True
            for sub_question in sub_questions:
                cells = table.add_row().cells
                cells[1].merge(cells[2])
                cells[1].text = "%s (n=%s)" % (sub_question.prompt, sub_question.n)
                for response in sub_question.responses:
                    if not response.frequencies:
                        shading_elm = parse_xml(r'<w:shd {} w:fill="FFF206"/>'.format(nsdecls('w')))
                        cells[1]._tc.get_or_add_tcPr().append(shading_elm)
                    for year, frequency in response.frequencies.items():
                        cells[3].text = self.avgs_percent(frequency, first_row)
                first_row = False

    def write_trended_allocate(self, sub_questions):
        headers = self.max_years_subquestions(sub_questions)
        table = self.doc.add_table(rows=1, cols=len(headers)+4)
        titles_row = table.add_row().cells
        titles_row[1].merge(titles_row[2])
        headers_index = 0
        while headers_index < len(headers):
            header_text = "Total %s" % headers[headers_index]
            titles_row[headers_index+4].text = header_text
            headers_index += 1
        first_row = True
        for sub_question in sub_questions:
            for response in sub_question.responses:
                region_cells = table.add_row().cells
                region_cells[1].merge(region_cells[3])
                region_cells[1].text = "%s (n=%s)" % (sub_question.prompt, sub_question.n)
                freq_col = 4
                if not response.frequencies:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="FFF206"/>'.format(nsdecls('w')))
                    region_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
                for header in headers:
                    if response.frequencies.get(header) is not None:
                        freq = response.frequencies.get(header)
                        text = self.avgs_percent(freq, first_row)
                        region_cells[freq_col].text = text
                    else:
                        if first_row is True:
                            region_cells[freq_col].text = "$--"
                        else:
                            region_cells[freq_col].text = "--"
                    first_row = False
                    freq_col += 1

    def write_matrix(self, sub_questions):
        if len(self.years) > 0:
            self.write_trended_matrix(sub_questions)
        else:
            table = self.doc.add_table(rows = 1, cols = 0)
            table.add_column(width = Inches(1))
            table.add_column(width = Inches(1))
            first_row = True
            header_cells = table.add_row().cells
            for sub_question in sub_questions:
                if first_row == True:
                    for response in sub_question.responses:
                        response_cells = table.add_column(width = Inches(1)).cells
                        response_cells[1].text = response.response
                question_cells = table.add_row().cells
                question_cells[1].text = "%s (n=%s)" % (sub_question.prompt, sub_question.n)
                index = 2
                for response in sub_question.responses:
                    if response.has_frequency is True:
                        for year, frequency in response.frequencies.items():
                            question_cells[index].text = self.freqs_percent(frequency, first_row)
                    else:
                        if not response.frequencies:
                            shading_elm = parse_xml(r'<w:shd {} w:fill="FFF206"/>'.format(nsdecls('w')))
                            question_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
                        if first_row is True:
                            question_cells[index].text = "--%"
                        else:
                            question_cells[index].text = "--"
                    first_row = False
                    index += 1

    def write_trended_matrix(self, sub_questions):
        self.doc.add_paragraph("") # space between questions
        for sub_question in sub_questions:
            paragraph = self.doc.add_paragraph() # each question starts a new paragraph
            self.write_name(sub_question.name, paragraph)
            self.write_prompt(sub_question.prompt, paragraph)
            self.write_n(sub_question.n, paragraph)
            self.write_responses(sub_question.responses, sub_question.stat)
            self.doc.add_paragraph("") # space between questions

    def max_years(self, responses):
        years_used = []
        for response in responses:
            for year in self.years:
                if response.frequencies.get(year) is not None:
                    if year not in years_used:
                        years_used.append(year)
        return years_used

    def max_years_subquestions(self, sub_questions):
        years_used = []
        for question in sub_questions:
            for response in question.responses:
                for year in self.years:
                    if response.frequencies.get(year) is not None:
                        if year not in years_used:
                            years_used.append(year)
        return years_used

    def avg_float(self, average, is_first):
        average = float(average)
        if average >= 0 and average < 1:
            result = '<1'
        else:
            result = average
        if is_first is True:
            result = str(result)
        return str(result)

    def freqs_percent(self, freq, is_first):
        result = 0
        if float(freq) >= 1.0:
            result = int(freq) * 100
        else:
            percent = float(freq)
            percent = percent * 100
            if percent >= 0 and percent < 1:
                result = "<1"
            else:
                result = int(round(percent))
        if is_first:
            result = str(result) + "%"
        return str(result)

    def avgs_percent(self, average, is_first):
        average = float(average)
        if average >= 0 and average < 1:
            result =  '<1'
        else:
            result = int(round(average))
        if is_first is True:
            result = "$" + str(result)
        return str(result)

