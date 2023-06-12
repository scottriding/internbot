import docx
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from collections import OrderedDict

class Document(object):

	def build_document_report(self, question_blocks, groups, path_to_template, path_to_output):
		self.__question_blocks = question_blocks
		self.__groups = groups
		self.__document = docx.Document(path_to_template)
		self.line_break = self.__document.styles['LineBreak']
		self.write_questions()
		self.__document.save(path_to_output)

	def write_questions(self):
		for question in self.__question_blocks.questions:
			if question.parent == "CompositeQuestion":
				if question.type == "CompositeMatrix":
					self.write_matrix(question)
				else:
					self.write_binary(question)
			elif question.type == "TE":
				self.write_openend(question)
			elif question.type == "DB":
				self.write_descriptive(question)
			else:
				self.write_question(question)
			
			line_break = self.__document.add_paragraph("")
			line_break.style = self.line_break
			self.__document.add_paragraph("")

	def write_matrix(self, matrix):
		paragraph = self.__document.add_paragraph()
		
		if len(self.__groups) == 0:
			paragraph.add_run(matrix.name + ".")
			self.write_prompt(matrix.prompt, paragraph)
			paragraph.add_run(matrix.base_n)
			
			table = self.__document.add_table(rows = 1, cols = 0)
			table.add_column(width = Inches(1))
			table.add_column(width = Inches(1))
			
			header_cells = table.add_row().cells
			is_header = True
			for question in matrix.questions:
				if is_header:
					for response in question.responses:
						response_cells = table.add_column(width = Inches(1)).cells
						response_cells[1].text = response.label
					is_header = False
				
				statement_cells = table.add_row().cells
				statement_cells[1].text = question.prompt
				index = 2
				first = True
				for response in question.responses:
					for group, frequency in response.frequencies.frequencies.items():
						if int(frequency.population) == 0:
							statement_cells[index].text = "*"
						elif frequency.stat == 'percent':
							percent = float(frequency.result) * 100
							if percent >= 0 and percent < 1:
								statement_cells[index].text = "<1"
							else:
								statement_cells[index].text = str(int(round(percent)))
							
							if first:
								statement_cells[index].text += "%"
								first = False
						else:
							statement_cells[index].text = frequency.result
						index += 1 
			
		else:
			paragraph_format = paragraph.paragraph_format
			paragraph_format.keep_together = True
			paragraph.add_run(matrix.prompt + ". ")
			paragraph.add_run(matrix.base_n)
			
			self.__document.add_paragraph("")
			self.__document.add_paragraph("")
			
			for question in matrix.questions:
				paragraph = self.__document.add_paragraph()
				paragraph.add_run(question.name + ".")
				self.write_prompt(question.prompt, paragraph)
				self.write_grouped_results(question.responses)
				
				self.__document.add_paragraph("")
				self.__document.add_paragraph("")
				

	def write_binary(self, binary):
		paragraph = self.__document.add_paragraph()
		paragraph.add_run(binary.name + ".")
		self.write_prompt(binary.prompt, paragraph)
		paragraph.add_run(binary.base_n)
		
		if len(self.__groups) == 0:
			table = self.__document.add_table(rows=1, cols=5)
			first = True
			cells = table.add_row().cells
			cells[1].merge(cells[2])
			cells[3].text = "Average"
			
			for question in binary.questions:
				question_cells = table.add_row().cells
				question_cells[1].merge(question_cells[2])
				for response in question.responses:
					for group, frequency in response.frequencies.frequencies.items():
						question_cells[1].text = question.prompt
						if int(frequency.population) == 0:
							question_cells[3].text = "*"
						elif frequency.stat == 'percent':
							percent = float(frequency.result) * 100
							if percent >= 0 and percent < 1:
								question_cells[3].text = "<1"
							else:
								question_cells[3].text = str(int(round(percent)))
							
							if first:
								question_cells[3].text += "%"
								first = False
						else:
							question_cells[3].text = frequency.result
		else:
			table = self.__document.add_table(rows=1, cols=len(self.__groups)+4)
			titles_row = table.add_row().cells
			titles_row[1].merge(titles_row[2])
			headers_index = 0
			first_row = True
			first_group = True
			row = -1
			table_rows = []
			while headers_index < len(self.__groups):
				header_text = "Average %s" % self.__groups[headers_index]
				titles_row[headers_index+4].text = header_text
				for sub_question in binary.questions:
					if first_group:
						cells = table.add_row().cells
						cells[1].merge(cells[2])
						cells[1].text = sub_question.prompt
						table_rows.append(cells)
						row += 1
					else:
						row += 1
					for response in sub_question.responses:
						for group, frequency in response.frequencies.frequencies.items():
							if group == self.__groups[headers_index]:
								if int(frequency.population) == 0:
									table_rows[row][headers_index+4].text = "*"
								elif frequency.stat == 'percent':
									percent = float(frequency.result) * 100
									if percent >= 0 and percent < 1:
										table_rows[row][headers_index+4].text = "<1"
									else:
										table_rows[row][headers_index+4].text = str(int(round(percent)))

									if first_row:
										table_rows[row][headers_index+4].text += "%"
										first_row = False
								else:
									table_rows[row][headers_index+4].text = frequency.result
				headers_index += 1
				row = -1
				first_group = False
			

	def write_openend(self, openend):
		paragraph = self.__document.add_paragraph()
		paragraph.add_run(openend.name + ".")
		self.write_prompt(openend.prompt, paragraph)
		if len(openend.responses) > 0:
			paragraph.add_run(openend.base_n)
			if len(self.__groups) == 0:
				self.write_results(openend.responses)
			else:
				self.write_grouped_results(openend.responses)
		else:
			paragraph.add_run(' (OPEN-ENDED RESPONSES VERBATIM IN APPENDIX)')

	def write_descriptive(self, description):
		paragraph = self.__document.add_paragraph()
		paragraph.add_run(description.prompt)

	def write_question(self, question):
		paragraph = self.__document.add_paragraph()
		paragraph.add_run(question.name + ".")
		self.write_prompt(question.prompt, paragraph)
		paragraph.add_run(question.base_n)
		if len(self.__groups) == 0:
			self.write_results(question.responses)
		else:
			self.write_grouped_results(question.responses)

	def write_prompt(self, prompt, paragraph):
		paragraph_format = paragraph.paragraph_format
		paragraph_format.keep_together = True
		paragraph_format.left_indent = Inches(1)
		paragraph.add_run("\t" + prompt + " ")
		paragraph_format.first_line_indent = Inches(-1)

	def write_results(self, responses):
		table = self.__document.add_table(rows=1, cols=5)
		first_row = True
		for response in responses:
			response_cells = table.add_row().cells
			response_cells[1].merge(response_cells[2])
			response_cells[1].text = response.label
			for group, frequency in response.frequencies.frequencies.items():
				if int(frequency.population) == 0:
					response_cells[3].text = "*"
				elif frequency.stat == 'percent':
					percent = float(frequency.result) * 100
					if percent >= 0 and percent < 1:
						response_cells[3].text = "<1"
					else:
						response_cells[3].text = str(int(round(percent)))

					if first_row:
						response_cells[3].text += "%"
						first_row = False
				else:
					response_cells[3].text = frequency.result

	def write_grouped_results(self, responses):
		table = self.__document.add_table(rows=1, cols=len(self.__groups) + 4)
		titles_row = table.add_row().cells
		titles_row[1].merge(titles_row[2])
		for header_index in range(len(self.__groups)):
			titles_row[header_index + 4].text = self.__groups[header_index]
		first_row = True
		for response in responses:
			response_cells = table.add_row().cells
			response_cells[1].merge(response_cells[3])
			response_cells[1].text = response.label
			freq_col = 4
			for group, frequency in response.frequencies.frequencies.items():
				if int(frequency.population) == 0:
					response_cells[freq_col].text = "*"
				elif frequency.stat == 'percent':
					percent = float(frequency.result) * 100
					if percent >= 0 and percent < 1:
						response_cells[freq_col].text = "<1"
					else:
						response_cells[freq_col].text = str(int(round(percent)))

					if first_row:
						response_cells[freq_col].text += "%"
						first_row = False
				else:
					response_cells[freq_col].text = frequency.result
				freq_col += 1   