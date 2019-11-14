import docx

class Document(object):

    def __init__ (self, path_to_template):
        self.__doc = docx.Document(path_to_template)

    def write_appendix(self, questions):
        first_question = True
        for question, value in questions.items():
            if first_question is False:
                self.__doc.add_page_break()
            paragraph = self.__doc.add_paragraph()
            self.write_question(value, paragraph)
            self.__doc.add_paragraph()
            first_question = False

    def write_question(self, question, paragraph):
        to_print = "Writing: %s" % question.name
        print(to_print)
        paragraph.add_run(question.name + ".")
        paragraph_format = paragraph.paragraph_format
        paragraph_format.keep_together = True
        paragraph_format.left_indent = docx.shared.Inches(1)
        prompt_to_add = "\t%s (n=%s)\n" % (question.prompt, question.response_count)
        paragraph.add_run(prompt_to_add)
        paragraph_format.first_line_indent = docx.shared.Inches(-1)

        self.write_responses(question.responses, paragraph)

    def write_responses(self, responses, paragraph):
        table = self.__doc.add_table(rows = 0, cols = 5)
        try:
            table.style = 'Appendix' # Custom format in template
        except KeyError:
            pass
        for response in responses:
            response_cells = table.add_row().cells
            response_cells[0].merge(response_cells[4])
            response_cells[0].text = response

    def save(self, path_to_output):
        self.__doc.save(path_to_output)
