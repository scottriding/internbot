from docx import Document
from docx.shared import Inches

class ToplineReport(object):

    def __init__ (self, questions, path_to_template):
        self.doc = Document(path_to_template)
        self.questions = questions

    def save(self, path_to_output):
        self.write_questions()
        self.save_file(path_to_output)

    def write_questions(self):
        for question in self.questions:
            if question.type == 'CompositeMatrix' or \
               question.type == 'CompositeMultipleSelect' or \
               question.type == 'CompositeHotSpot' or \
               question.type == 'CompositeConstantSum':
                self.write_composite_question(question)
            elif question.type == 'TE':
                self.write_open_ended(question)
            else:
                self.write_question(question)

    def save_file(self, path_to_output):
        self.doc.save(path_to_output)

    def write_question(self, question):
        paragraph = self.doc.add_paragraph() # each question starts a new paragraph
        self.write_name(question.name, paragraph)
        self.write_prompt(question.prompt, paragraph)
        self.write_responses(question.responses)
        self.doc.add_paragraph("") # space between questions

    def write_composite_question(self, question):
        paragraph = self.doc.add_paragraph() # each question starts a new paragraph
        self.write_name(question.name, paragraph)
        self.write_prompt(question.prompt, paragraph)
        if question.type == 'CompositeMatrix':
            self.write_matrix(question.questions)
        elif question.type == 'CompositeConstantSum':
            self.write_allocate(question.questions)
        else:
            self.write_binary(question.questions)
        self.doc.add_paragraph("") # space between questions

    def write_open_ended(self, question):
        paragraph = self.doc.add_paragraph() # each question starts a new paragraph
        self.write_name(question.name, paragraph)
        self.write_prompt(question.prompt, paragraph)
        paragraph.add_run(' (OPEN-ENDED RESPONSES VERBATIM IN APPENDIX)')
        self.doc.add_paragraph("") # space between questions

    def write_name(self, name, paragraph):
        paragraph.add_run(name + ".")

    def write_prompt(self, prompt, paragraph):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.keep_together = True # question prompt will all be fit in one page
        paragraph_format.left_indent = Inches(1)
        paragraph.add_run("\t" + prompt)
        paragraph_format.first_line_indent = Inches(-1) # hanging indent if necessary

    def write_n(self, n, paragraph):
        paragraph.add_run("(n = " + str(n) + ")")

    def write_responses(self, responses):
        table = self.doc.add_table(rows = 1, cols = 5)
        first_row = True
        for response in responses:
            if response.type == 'Response' or response.type == 'NAResponse':
                response_cells = table.add_row().cells
                response_cells[1].merge(response_cells[2])
                if response.type == 'NAResponse':
                    response_cells[1].text = 'Don\'t know / Skipped'
                else:
                    response_cells[1].text = response.response
                if response.has_frequency is True and first_row is True:
                    response_cells[3].text = self.freqs_percent(response.frequency) + "%"
                    first_row = False
                elif response.has_frequency is True and first_row is False:
                    response_cells[3].text = self.freqs_percent(response.frequency)
                elif response.has_frequency is False and first_row is True:
                    response_cells[3].text = '--%'
                    first_row = False
                else:
                    response_cells[3].text = '--'

    def write_binary(self, sub_questions):
        table = self.doc.add_table(rows = 1, cols = 5)
        first_row = True
        for sub_question in sub_questions:
            region_cells = table.add_row().cells
            region_cells[1].merge(region_cells[2])
            response = next((response for response in sub_question.responses if response.response == '1'), None)
            region_cells[1].text = sub_question.prompt
            if response.has_frequency is True and first_row is True:
                region_cells[3].text = self.freqs_percent(response.frequency) + "%"
                first_row = False
            elif response.has_frequency is True and first_row is False:
                region_cells[3].text = self.freqs_percent(response.frequency)
            elif response.has_frequency is False and first_row is True:
                region_cells[3].text = '--%'
                first_row = False
            else:
                region_cells[3].text = '--'

    def write_allocate(self, sub_questions):
        table = self.doc.add_table(rows = 1, cols = 5)
        first_row = True
        for sub_question in sub_questions:
            region_cells = table.add_row().cells
            region_cells[1].merge(region_cells[2])
            for response in sub_question.responses:
                region_cells[1].text = sub_question.prompt
                if first_row is True and response.has_frequency is True:
                    region_cells[3].text = '$%s' % self.avgs_percent(response.frequency)
                    first_row = False
                elif first_row is True and response.has_frequency is False:
                    region_cells[3].text = '$--'
                    first_row = False
                elif first_row is False and response.has_frequency is True:
                    region_cells[3].text = self.avgs_percent(response.frequency)
                else:
                    region_cells[3].text = '--'

    def write_matrix(self, sub_questions):
        table = self.doc.add_table(rows = 1, cols = 0)
        table.add_column(width = Inches(1))
        table.add_column(width = Inches(1))
        first_row = True
        header_cells = table.add_row().cells
        for sub_question in sub_questions:
            if first_row == True:
                for response in sub_question.responses:
                    response_cells = table.add_column(width = Inches(1)).cells
                    response_cells[1].text = response.response
            question_cells = table.add_row().cells
            question_cells[1].text = sub_question.prompt
            index = 2
            for response in sub_question.responses:
                if response.has_frequency is True and first_row is True:
                    question_cells[index].text = self.freqs_percent(response.frequency) + '%'
                    first_row = False
                elif response.has_frequency is True and first_row is False:
                    question_cells[index].text = self.freqs_percent(response.frequency)
                elif response.has_frequency is False and first_row is True:
                    question_cells[index].text = '--%'
                    first_row = False
                else:
                    question_cells[index].text = '--'
                index += 1

    def get_doc(self):
        return self.doc

    def freqs_percent(self, freq):
        percent = freq * 100
        if percent > 0 and percent < 1:
            return "<1"
        elif percent == 0:
            return "*"
        result = int(round(percent))
        return str(result)

    def avgs_percent(self, average):
        if average > 0 and average < 1:
            return '<1'
        elif average == 0:
            return '*'
        result = int(round(average))
        return str(result)
