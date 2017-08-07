from docx import Document
from docx.shared import Inches

class ToplineAppendix(object):

    def __init__(self):
        doc = Document()

    def write_independent(self, questions, path_to_template):
        self.doc = Document(path_to_template)
        self.write_appendix(questions)

    def write_with_topline(self, questions, path_to_topline):
        self.doc = Document(path_to_topline)
        self.write_title()
        self.write_appendix(questions)

    def save(self, path_to_output):
        self.doc.save(path_to_output)
    
    def write_appendix(self, questions):
        first_question = True
        for question in questions:
            if first_question is False:
                self.doc.add_page_break()
            paragraph = self.doc.add_paragraph()
            self.write_question(question, paragraph)
            self.write_responses(question.responses, paragraph)
            self.doc.add_paragraph()
            first_question = False

    def write_title(self):
        self.doc.add_page_break()
        self.doc.add_heading('Appendix')
        self.doc.add_paragraph()

    def write_question(self, question, paragraph):
        paragraph.add_run(question.name + ".")
        paragraph_format = paragraph.paragraph_format
        paragraph_format.keep_together = True # question prompt will all be fit in one page
        paragraph_format.left_indent = Inches(1)
        paragraph.add_run("\t" + question.prompt + "\n")
        paragraph_format.first_line_indent = Inches(-1) # hanging indent if necessary

    def write_responses(self, responses, paragraph):
        table = self.doc.add_table(rows = 0, cols = 5)
        for response in responses:
            response_cells = table.add_row().cells
            response_cells[0].merge(response_cells[4])
            try:
                response_cells[0].text = response.response
            except ValueError:
                pass