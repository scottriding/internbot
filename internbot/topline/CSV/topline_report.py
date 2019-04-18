from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

class ToplineReport(object):

    def __init__(self, questions, path_to_template, years):
        self.doc = Document(path_to_template)
        self.line_break = self.doc.styles['LineBreak']
        self.questions = questions
        self.headers = years
        names = questions.list_names()
        for name in names:
            self.write_question(name)
        print "Finished!"
        
    def write_question(self, name):
        question = self.questions.get(name)
        print "Writing question: %s" % name
        if question.display_logic != "":
            display_prompt = self.doc.add_paragraph()
            display_prompt.add_run("(" + question.display_logic + ")")
            self.doc.add_paragraph()
        paragraph = self.doc.add_paragraph() # each question starts a new paragraph
        self.write_name(name, paragraph)
        self.write_prompt(question.prompt, paragraph)
        self.write_n(question.n, paragraph)
        if len(question.responses) > 0:
            self.write_responses(question.responses)
        new = self.doc.add_paragraph("") # space between questions
        new.style = self.line_break
        self.doc.add_paragraph("") # space between questions
        
    def write_name(self, name, paragraph):
        paragraph.add_run(name + ".")
    
    def write_prompt(self, prompt, paragraph):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.keep_together = True           # question prompt will all be fit in one page
        paragraph_format.left_indent = Inches(1)        
        paragraph.add_run("\t" + prompt)                
        paragraph_format.first_line_indent = Inches(-1) # hanging indent if necessary
        
    def write_n(self, n, paragraph):
        paragraph.add_run(" (n = " + str(n) + ")")
    
    def write_responses(self, responses):
        first_response_freqs = responses[0].frequencies
        data_columns  = len(first_response_freqs)
        table = self.doc.add_table(rows=1, cols=data_columns+5)
        if data_columns > 1:
            titles_row = table.add_row().cells  
            headers_index = 0
            for i in range(5, data_columns+5):
                header_text = "Total %s" % self.headers[headers_index]
                titles_row[i].text = header_text
                headers_index += 1
        first_row = True
        first_freq = True
        for response in responses:
            response_cells = table.add_row().cells
            response_cells[2].merge(response_cells[3])
            response_cells[2].text = response.name
            freq_col = 5
            for freq in response.frequencies:
                text = self.freqs_percent(freq, first_freq)
                if first_freq is True:
                	first_freq = False
                response_cells[freq_col].text = text
                freq_col += 1

    def save(self, path_to_output):
        self.doc.save(path_to_output)
        
    def freqs_percent(self, freq, is_first=False):
        result = 0
        if float(freq) >= 1.0:
            result = int(freq) * 100
        else:
            percent = float(freq) * 100
            if percent >= 0 and percent < 1:
                return "<1"
            result = int(round(percent))
        if is_first:
            result = str(result) +"%"
        return str(result)