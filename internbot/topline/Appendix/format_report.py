from openpyxl import load_workbook, Workbook
from openpyxl.styles.borders import Border, Side
from openpyxl.styles import PatternFill, Font, Alignment
from openpyxl.drawing.image import Image
from docx import Document
from docx.shared import Inches
from collections import OrderedDict

class SSAppendixBuilder(object):

    def __init__(self, is_qualtrics):
        self.__workbook = Workbook()
        self.__is_qualtrics = is_qualtrics

        if is_qualtrics:
            # fill colors
            self.__header_fill = PatternFill("solid", fgColor = "1E262E")
        else:
            self.__header_fill = PatternFill("solid", fgColor = "0F243E")

        self.__table_fill = PatternFill("solid", fgColor = "E7E6E6")
       
        # font styles
        self.__font_reg = Font(name = 'Arial', size = 8)
        self.__font_bold = Font(name = 'Arial', size = 8, bold = True)
        self.__font_back_hyperlink = Font(name = 'Arial', size = 8, color = "A2AAAD")

        # alignments
        self.__align_center = Alignment(horizontal = "center", vertical = "center", wrapText = True)
        self.__align_left = Alignment(horizontal = "left", vertical = "center", wrapText = True)

        # borders
        self.__thin_bottom = Border(bottom = Side(style = 'thin'))

    def write_appendix(self, questions):
        if self.__workbook.get_sheet_by_name("Sheet") is not None:
            toc_sheet = self.__workbook.get_sheet_by_name('Sheet')
            toc_sheet.title = "TOC"
        else:
            toc_sheet = self.__workbook.create_sheet("TOC")

        self.write_toc(toc_sheet, questions)

        for question, value in questions.iteritems():
            new_sheet = self.__workbook.create_sheet(question)
            self.write_question(value, new_sheet)

    def write_toc(self, sheet, questions):
        print "Writing TOC"
        self.write_header(sheet)

        sheet["A2"].value = "Question Name"
        sheet["A2"].font = self.__font_bold
        sheet["A2"].alignment = self.__align_center

        sheet["B2"].value = "Question Title"
        sheet["B2"].font = self.__font_bold
        sheet["B2"].alignment = self.__align_center

        sheet["C2"].value = "Verbatim response count"
        sheet["C2"].font = self.__font_bold
        sheet["C2"].alignment = self.__align_center

        current_row = 3
        for question, value in questions.iteritems():
            question_name = "A%s" % str(current_row)
            question_title = "B%s" % str(current_row)
            base_size = "C%s" % str(current_row)

            sheet[question_name].value = "=HYPERLINK(\"#'%s'!A1\",\"%s\")" % (question, question)
            sheet[question_name].font = Font(name = 'Arial', size = 8, underline = 'single')
            sheet[question_name].alignment = self.__align_center
            sheet[question_name].border = self.__thin_bottom

            sheet[question_title].value = value.prompt
            sheet[question_title].font = self.__font_reg
            sheet[question_title].alignment = self.__align_center
            sheet[question_title].border = self.__thin_bottom

            sheet[base_size].value = value.response_count
            sheet[base_size].font = self.__font_reg
            sheet[base_size].alignment = self.__align_center
            sheet[base_size].border = self.__thin_bottom
            current_row += 1

    def write_header(self, sheet):
        sheet.column_dimensions["A"].width = 10
        sheet.column_dimensions["B"].width = 55
        sheet.column_dimensions["C"].width = 33

        if self.__is_qualtrics:
            sheet.row_dimensions[1].height = 35
        else:
            sheet.row_dimensions[1].height = 52

        sheet["A1"].fill = self.__header_fill
        sheet["B1"].fill = self.__header_fill
        sheet["C1"].fill = self.__header_fill

        if self.__is_qualtrics:
            logo = Image("templates_images/QLogo.png")
        else:
            logo = Image("templates_images/y2_xtabs.png")
        sheet.add_image(logo, "C1")

    def write_question(self, question, sheet):
        print "Writing: %s" % question.name 
        self.write_header(sheet)
        sheet.row_dimensions[2].height = 36

        sheet["A1"].value = '=HYPERLINK("#TOC!A1","Return to Table of Contents")'
        sheet["A1"].font = self.__font_back_hyperlink
        sheet["A1"].alignment = self.__align_center

        sheet.merge_cells(start_column=1, end_column=3, start_row=2, end_row=2)

        sheet["A2"].fill = self.__table_fill
        sheet["B2"].fill = self.__table_fill
        sheet["C2"].fill = self.__table_fill

        sheet["A2"].value = "%s: %s" % (question.name, question.prompt)
        sheet["A2"].alignment = self.__align_center
        sheet["A2"].font = self.__font_reg

        self.write_responses(question.responses, sheet)

    def write_responses(self, responses, sheet):
        current_row = 3
        for response in responses:
            sheet.row_dimensions[current_row].height = 25
            a_cell = "A%s" % str(current_row)
            b_cell = "B%s" % str(current_row)
            c_cell = "C%s" % str(current_row)
            sheet.merge_cells(start_row=current_row, end_row=current_row, start_column = 1, end_column = 3)

            sheet[a_cell].value = response
            sheet[a_cell].font = self.__font_reg
            sheet[a_cell].alignment = self.__align_left
            sheet[a_cell].border = self.__thin_bottom
            sheet[b_cell].border = self.__thin_bottom
            sheet[c_cell].border = self.__thin_bottom

            current_row += 1

    def save(self, path_to_output):
        self.__workbook.save(path_to_output)

class DocAppendixBuilder(object):

    def __init__ (self, path_to_template, is_qualtrics=False):
        self.__doc = Document(path_to_template)

    def write_appendix(self, questions):
        first_question = True
        for question, value in questions.iteritems():
            if first_question is False:
                self.__doc.add_page_break()
            paragraph = self.__doc.add_paragraph()
            self.write_question(value, paragraph)
            self.__doc.add_paragraph()
            first_question = False

    def write_question(self, question, paragraph):
        print "Writing: %s" % question.name
        paragraph.add_run(question.name + ".")
        paragraph_format = paragraph.paragraph_format
        paragraph_format.keep_together = True
        paragraph_format.left_indent = Inches(1)
        prompt_to_add = "\t%s (n=%s)\n" % (question.prompt, question.response_count)
        paragraph.add_run(prompt_to_add)
        paragraph_format.first_line_indent = Inches(-1)

        self.write_responses(question.responses, paragraph)

    def write_responses(self, responses, paragraph):
        table = self.__doc.add_table(rows = 0, cols = 5)
        for response in responses:
            response_cells = table.add_row().cells
            response_cells[0].merge(response_cells[4])
            response_cells[0].text = response

    def save(self, path_to_output):
        self.__doc.save(path_to_output)

class OpenEndQuestion(object):

    def __init__(self, name, prompt):
        self.__name = name
        self.__prompt = prompt
        self.__responses = []
        self.__response_count = 0

    def add_response(self, response):
        if response != "":
            self.__responses.append(response)
            self.__response_count += 1
            self.__responses.sort()

    @property
    def name(self):
        return self.__name

    @property
    def prompt(self):
        return self.__prompt

    @property
    def responses(self):
        return self.__responses

    @property
    def response_count(self):
        return self.__response_count